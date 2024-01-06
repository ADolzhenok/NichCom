from voice_assistant.get_audio import get_audio_rus
from voice_assistant.speak import speak_rus
import os
from time import strftime, sleep
import docx
import pyautogui as pg
import winshell
import matplotlib.pyplot as plt
import requests
# import pyodbc
from voice_assistant.voice_assistant_RUS import main_rus


def timer():
    speak_rus("Хорошо, я включаю таймер")
    speak_rus("О чет Вы хотите, чтобы я напомнила?")
    reminder_name = get_audio_rus().lower()
    speak_rus("Через сколько минут?")
    flag = 0
    while flag == 0:
        digit = get_audio_rus()
        for char in digit:
            if char.isdigit():
                digit = char
                break
        try:
            minutes = int(digit)
            minutes = minutes * 60
            flag += 1
            sleep(minutes)
            outcome = speak_rus(f"Вам следует {reminder_name}")
            # mongo_collect.insert_one(
            #     {"user": user, "action": "Включение таймера",
            #      "details": [
            #          {"[ACTION]": "ВКЛЮЧИТЬ ТАЙМЕР"},
            #          {"[TIME]": f'{strftime("%c")}'},
            #          {"[REMIND ABOUT]": f'{reminder_name}'},
            #          {"[HOW MANY MINUTES]": f'{minutes}'}
            #      ],
            #      "status": "SUCCESS",
            #      "outcome": f"{outcome }"
            #      }
            # )
        except ValueError:
            outcome = speak_rus("Пожалуйста, назовите число")
            # mongo_collect.insert_one(
            #     {"user": user, "action": "Turn on the timer",
            #      "details": [
            #          {"[ACTION]": "TIMER ON"},
            #          {"[TIME]": f'{strftime("%c")}'},
            #          {"[REMIND ABOUT]": f"{reminder_name}"},
            #          {"[HOW MANY MINUTES]": f"{minutes}"}
            #      ],
            #      "status": "FAIL",
            #      "outcome": f"{outcome}"
            #      }
            # )


def time():
    speak_rus(f'Сейчас {strftime("%X")}')
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Озвучивание текущего времени",
    #      "details": [
    #          {"[ACTION]": "ТЕКУЩЕЕ ВРЕМЯ"},
    #          {"[TIME]": f'{strftime("%c")}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f"{outcome}"
    #      }
    # )


def money():
    city = 'Гомель'
    api_not_json = requests.get(f'https://belarusbank.by/api/kursExchange?city={city}')
    api = api_not_json.json()
    speak_rus("Какую валюту вы хотите получить?")
    currency = get_audio_rus()
    speak_rus("Вы хотите продать или купить валюту?")
    transaction = get_audio_rus()
    if 'доллар' or '$' in currency:
        currency = 'USD'
    elif 'евро' in currency:
        currency = 'EUR'
    elif 'рубль' in currency:
        currency = 'RUB'
    if 'продать' in transaction:
        value = api[0][f'{currency}_in']
        outcome = speak_rus(f"Вы можете продать {currency} за {value}")
    else:
        value = api[0][f'{currency}_out']
        outcome = speak_rus(f"Вы можете купить {currency} за {value}")
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Получение информации о валютах",
    #      "details": [
    #          {"[ACTION]": "ИНФО ВАЛЮТЫ "},
    #          {"[TIME]": f'{strftime("%c")}'},
    #          {"[CURRENCY]": f'{currency}'},
    #          {"[TRANSACTION]": f'{transaction}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f"{outcome}"
    #      }
    # )


def weather():
    appid = "66eb8cf1abebae782e1d3700a75e203e"

    def form_url_string(s_request):
        nonlocal appid
        s_appid = "&APPID=" + appid
        s_template = "http://api.openweathermap.org/data/2.5/" + s_request + s_appid
        return s_template

    s_city = "Gomel"
    s_country = "BY"
    s_request = "weather?q={},{}&units=metric".format(s_city, s_country)
    s_search_url = form_url_string(s_request)
    try:
        res = requests.get(s_search_url)
        data = res.json()
        outcome = speak_rus(
            f"Погода сейчас:\n Температура: {data['main']['temp']}")
    except Exception as e:
        speak_rus("Извините, город не найден")
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Озвучивание текущей погоды",
    #      "details": [
    #          {"[ACTION]": "ТЕКУЩАЯ ПОГОДА"},
    #          {"[TIME]": f'{strftime("%c")}'},
    #          {"[CITY]": f'{s_city}'},
    #          {"[COUNTRY]": f'{s_country}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f'{outcome}'
    #      }
    # )


def creating_word_file():
    speak_rus('Скажите имя файла')
    file_name = get_audio_rus()
    document = docx.Document()
    speak_rus(f'Имя файла : {file_name}')
    speak_rus('Создать пустой файл?')
    accept = get_audio_rus()
    list_words = ('да', 'ага', 'создать пустой')
    for i in list_words:
        if i in accept:
            document.save(f'{file_name}.docx')
            os.system(f'{file_name}.docx')
            break
    else:
        speak_rus('Скажите текст, который будет в заголовке')
        title_word = get_audio_rus()
        first_letter = title_word[0].upper()
        document.add_heading(f'{first_letter}{title_word[1:]}', 0)
        speak_rus('Скажите текст, который будет внутри файла')
        just_text = get_audio_rus()
        first_letter = just_text[0].upper()
        p = document.add_paragraph(f'{first_letter}{just_text[1:]}  ')
        p.add_run(f'{first_letter}{just_text[1:]}  ').bold = True
        p.add_run(f'{first_letter}{just_text[1:]}.').italic = True
        document.save(f'{file_name}.docx')
        os.system(f'{file_name}.docx')
    outcome = speak_rus("Файл создан")
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Создание файла",
    #      "details": [
    #          {"[ACTION]": "СОЗДАТЬ ФАЙЛ"},
    #          {"[TIME]": f'{strftime("%c")}'},
    #          {"[FILE's NAME]": f'{file_name}'},
    #          {"[FILE's TYPE]": f'{""}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f'{outcome}'
    #      }
    # )


def open_my_computer():
    os.startfile(r"shell:mycomputerfolder")
    outcome = speak_rus('Мой Компьютер открыт')
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Открытие папки Мой Компьютер",
    #      "details": [
    #          {"[ACTION]": "ОТКРЫТЬ МОЙ КОМПЬЮТЕР"},
    #          {"[TIME]": f'{strftime("%c")}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f'{outcome}'
    #      }
    # )


def cleaning_recycle_bin():
    speak_rus("Вы действительно этого хотите?")
    text = get_audio_rus()
    if "да".lower() in text:
        outcome = speak_rus("Я очищаю")
        try:
            winshell.recycle_bin().empty(confirm=False, show_progress=False, sound=False)
        except BaseException:
            outcome = speak_rus("Ваша корзина пуста, нет необходимости в отчистке")
    elif "нет".lower() in text:
        outcome = speak_rus("Корзина не отчистилась")
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Очистка корзины",
    #      "details": [
    #          {"[ACTION]": "ОЧИСТИТЬ КОРЗИНУ"},
    #          {"[TIME]": f'{strftime("%c")}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f'{outcome}'
    #      }
    # )


def opening_the_browser():
    pg.hotkey('win', 'r')
    speak_rus("Открываю браузер")
    pg.typewrite('msedge')
    pg.press('enter')
    speak_rus("Что Вы хотите найти?")
    text = get_audio_rus()
    speak_rus(f"Я ищу : {text}")
    transf = text.maketrans("й ц у к е н г ш щ з х ъ ф ы в а п р о л д ж э я ч с м и т ь б ю", \
                            "q w e r t y u i o p [ ] a s d f g h j k l ; ' z x c v b n m , .")
    pg.typewrite(f'{text.translate(transf)}')
    print(text)
    pg.typewrite(["enter"])
    outcome = speak_rus(f"Я открыл браузер и ищу следующий текст:  {text}")
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Открытие браузера",
    #      "details": [
    #          {"[ACTION]": "ОТКРЫТЬ БРАУЗЕР"},
    #          {"[TIME]": f'{strftime("%c")}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f'{outcome}'
    #      }
    # )


def check_command():
    speak_rus("Какую команду Вы хотите проверить?")
    name_command = get_audio_rus()
    speak_rus("Сколько раз?")
    count = 0
    while count == 0:
        digit = get_audio_rus()
        for char in digit:
            if char.isdigit():
                digit = char
                break
        try:
            kol = int(digit)
            count = 1
        except ValueError:
            speak_rus("Извините, скажите число")
    x = []
    y = []
    for number in range(kol):
        if number != 0:
            speak_rus("Скажите имя команды")
            name_command = get_audio_rus()
        main_rus.main_rus_va(name_command)
        speak_rus('Как Вам результат, напечатайте 1, если результат Вас устроил, иначе 0')
        x.append(number)
        y.append(int(input()))
    fig = plt.figure()
    ax = fig.add_subplot(111)
    speak_rus("Скажите имя графика")
    name = input()
    ax.set(title=f'{name}')
    ax.set_xlabel('Количество вызова команды')
    ax.set_ylabel('Оценка за результат')
    plt.plot(x, y)
    plt.show()
    outcome = speak_rus("Я проверил комманды. Пожалуйста, посмотрите на экран. Я сделал график")
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Проверка команд",
    #      "details": [
    #          {"[ACTION]": "ПРОВЕРИТЬ КОМАНДЫ"},
    #          {"[TIME]": f'{strftime("%c")}'},
    #          {"[NAME OF COMMAND]": f'{name_command}'},
    #          {"[HOW MANY TIMES CHECK]": f'{kol}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f'{outcome}'
    #      }
    # )


def calculator():
    speak_rus("Что Вы хотите вычислить?")
    text = get_audio_rus()
    text = text.split()
    try:
        if "+" in text[1] or "плюс" in text[1]:
            outcome = speak_rus(f'{text[0]} плюс {text[2]} равно {int(text[0]) + int(text[2])}')
        elif "-" in text[1] or "минус" in text[1]:
            outcome = speak_rus(f'{text[0]} минус {text[2]} равно {int(text[0]) - int(text[2])}')
        elif "*" in text[1] or "умножить" in text[1] or 'x' in text[1]:
            outcome = speak_rus(f'{text[0]} умножить на {text[2]} равно {int(text[0]) * int(text[2])}')
        elif "/" in text[1] or "разделить" in text[1] or "делить" in text[1]:
            outcome = speak_rus(f'{text[0]} разделить на {text[2]} равно {int(text[0]) / int(text[2])}')
    except ZeroDivisionError:
        outcome = speak_rus("Деление на 0 невозможно")
    except Exception:
        outcome = speak_rus("Извините, произошла ошибка")
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Вычисление",
    #      "details": [
    #          {"[ACTION]": "ВЫЧИСЛИТЬ ЗНАЧЕНИЯ"},
    #          {"[TIME]": f'{strftime("%c")}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f'{outcome}'
    #      }
    # )


def capability(mongo_collect, user):
    pass
    # conn = pyodbc.connect('Driver={SQL Server};'
    #                       'Server=USER-PC\SQLEXP;'
    #                       'Database=Voice_assistant;'
    #                       'Trusted_Connection=yes;')
    # cursor = conn.cursor()
    # try:
    #     cursor.execute('''drop table dbo.SKILLS_RUS''')
    # except pyodbc.ProgrammingError:
    #     pass
    # cursor.execute('''Create table dbo.SKILLS_RUS(
    #                                     Name varchar(60) unique
    #                                   )'''
    #                )
    # cursor.execute('''Insert into dbo.SKILLS_RUS values
    #                                     ('Включить таймер'),
    #                                     ('Произносить текущее время'),
    #                                     ('Произносить текущую погоду в Гомеле'),
    #                                     ('Создавать Вордовский файл'),
    #                                     ('Открывать папку "Мой компьютер'),
    #                                     ('Очищать корзину'),
    #                                     ('Открывать браузер и осуществлять в нем поиск'),
    #                                     ('Вычислять простые математические действия')
    #                                 ''')
    # cursor.execute('SELECT Name FROM Voice_assistant.dbo.SKILLS_RUS order by Name')
    # outcome = []
    # outcome.append(speak_rus("I can: "))
    # for row in cursor:
    #     outcome.append(speak_rus(row[0]))
    # outcome = '\n'.join(outcome)
    # conn.commit()
    # conn.close()
    # mongo_collect.insert_one(
    #     {"user": user, "action": "Озвучивание возможностей",
    #      "details": [
    #          {"[ACTION]": "ВОЗМОЖНОСТИ ГП"},
    #          {"[TIME]": f'{strftime("%c")}'}
    #      ],
    #      "status": "SUCCESS",
    #      "outcome": f'{outcome}'
    #      }
    # )


def search(mongo_collect, user):
    pg.hotkey('winleft', 's')
    speak_rus("Что Вы хотите найти?")
    file_name = get_audio_rus()
    transf = file_name.maketrans("й ц у к е н г ш щ з х ъ ф ы в а п р о л д ж э я ч с м и т ь б ю", \
                                 "q w e r t y u i o p [ ] a s d f g h j k l ; ' z x c v b n m , .")
    file_name = file_name.translate(transf)
    pg.typewrite(f'{file_name}')
    outcome = speak_rus("Посмотрите на экран, я закончил поиск")


def change_keyboard_lang(mongo_collect, user):
    pg.hotkey('win', 'space')
    outcome = speak_rus('Язык изменен')


def volume(action, mongo_collect, user):
    if action == "уменьшить":
        pg.press('volumedown')
    else:
        pg.press('volumeup')
    outcome = speak_rus("Я изменил громкость")



def task_manager(mongo_collect, user):
    outcome = speak_rus("Диспетчер задач открыт")
    pg.hotkey('ctrl', 'shift', 'esc')


# def history(name, cursor, mongo_collect):
#     speak_rus(f"Ты хочешь увидеть свою историю? "
#               f"Пожалуйста скажи да, если хочешь увидеть свою, нет если хочешь увидеть историю другого")
#     answer = get_audio_rus()
#     if "да" in answer:
#         cursor.execute(f"SELECT Действие FROM Voice_assistant.dbo.История where Имя = '{name}'")
#         for row in cursor:
#             speak_rus(row[0])
#     else:
#         speak_rus("Пожалуйста скажи полное имя пользователя, которого хочешь проверить")
#         another_name = get_audio_rus()
#         cursor.execute(f"SELECT Действие FROM Voice_assistant.dbo.История where Имя = '{another_name}'")
#         speak_rus(f"Следующие действия были осуществлены пользователем {name}")
#     outcome = []
#     outcome.append(speak_rus(f"История пользователя:  {name} следующая: "))
#     for row in cursor:
#         outcome.append(row[0])
#     outcome = '\n '.join(outcome)
#     print(type(outcome))
#     print(outcome)
#     mongo_collect.insert_one(
#         {"user": name, "action": "Отображение истории",
#          "details": [
#              {"[ACTION]": "ПОКАЗАТЬ ИСТОРИЮ"},
#              {"[TIME]": f'{strftime("%c")}'}
#          ],
#          "status": "SUCCESS",
#          "outcome": f"{outcome}"
#          }
#     )

